#!/usr/bin/env python3
"""PRAMANA 기술보고서 .docx 생성 — 처음 보는 사람이 깊게 이해(핵심만). 모델·작동·근거·정직한 한계·레퍼."""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import data
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

d=Document()
st=d.styles["Normal"].font; st.name="Calibri"; st.size=Pt(10.5)
def H(t,l=1,c=None):
    h=d.add_heading(t,level=l)
    if c:
        for r in h.runs: r.font.color.rgb=RGBColor(*c)
    return h
def P(t,b=False,i=False,sz=10.5,c=None):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=b; r.italic=i; r.font.size=Pt(sz)
    if c: r.font.color.rgb=RGBColor(*c)
    return p
def BUL(items):
    for it in items: d.add_paragraph(it,style="List Bullet")
def TBL(rows,hdr=True):
    t=d.add_table(rows=len(rows),cols=len(rows[0])); t.style="Light Grid Accent 1"
    for i,row in enumerate(rows):
        for j,v in enumerate(row):
            cell=t.cell(i,j); cell.text=str(v)
            if i==0 and hdr:
                for r in cell.paragraphs[0].runs: r.bold=True
    return t

# ── 표지 ──
ti=d.add_heading("PRAMANA — Systematic Equity/ETF Paper Trading System",0)
P("기술 보고서 v0.1 · 2026-06-11 · US-only · 가상자본 ₩100M (PAPER · NO LIVE)",i=True,c=(0x6b,0x72,0x80))
P("⚠️ 이 문서는 투자자문·종목추천·성과보장이 아니다. 전부 페이퍼(가상) 백테스트/시뮬이며 실거래 기록이 아니다.",b=True,c=(0x92,0x40,0x0e))

H("0. 한 줄 요약",1)
P("'좋은 종목을 찍는 AI'가 아니라, 데이터·비용 정직성과 사전등록 검증을 거쳐 가짜 알파를 걸러내고 — 살아남은 무상관 전략들을 결합·레버리지하는 솔로+AI용 systematic 트레이딩 OS다. "
  "현재 결과: 백테스트 공격버전 Sharpe ~1.1·연 +14%·MDD −17%(₩100M→3.37억). 단 수익은 ETF 추세 sleeve가 견인하며 강추세 레짐에 유리(forward 정직 기대 Sharpe ~0.5–0.6). 실거래는 12개월 forward 검증+사람 승인 전까지 금지.")

H("1. 목표와 철학",1)
BUL([
 "목표: 솔로/소규모+AI가 운영하는 research→validation→trading 운영체계. 단일 마법신호가 아님.",
 "순서 원칙(LOCKED): 데이터 정직 → 비용 정직 → 검증 통과 → 그 다음에야 모델 비교.",
 "사전등록 kill-gate: 결과 보기 전에 탈락조건을 박는다(goalpost 이동·과적합 방지).",
 "역할 분리(LOCKED): 신호≠주문, 알파≠포지션, 결정적 risk engine이 최종 veto, LLM은 off-path 자문만, 사람이 capital(실제 돈) 게이트.",
])

H("2. 시스템 아키텍처 (8 레이어)",1)
P("데이터 → 신호/모델 → alpha 앙상블 → 포트폴리오 → 리스크엔진 → 실행 → 모니터/kill → 거버넌스. 전부 config-driven·API-free(캐시 동결)·동결 재현(같은 입력=같은 숫자).",i=True)
TBL([
 ["레이어","모듈","핵심"],
 ["1 데이터","data.py","Sharadar 주식·ETF·LETF·vol · PIT·survivorship-safe · manifest/hash 무결성"],
 ["2 신호","features.py","동결 6 feature(value·quality·momentum·lowvol·event)+추세·vol · 정의 잠금"],
 ["3 앙상블","combined/portfolio_max/book_final","무상관 sleeve 위험예산 결합"],
 ["4 포트폴리오","(book_final)","위험예산 가중 → 목표 비중"],
 ["5 리스크엔진","risk_engine.py","결정적 veto·fractional-Kelly 사이징·4종 kill·gap 추정"],
 ["6 실행","execution.py","per-name 체결·스프레드/충격/차입/부분체결 비용"],
 ["7 모니터/kill","monitor.py","일일 지표·K1~6 kill·heartbeat·promotion gate"],
 ["8 거버넌스","(promotion gate)","paper→live는 사람 승인 게이트·LLM 비결정"],
])

H("3. 전략 — 무슨 '모델'을 왜, 어떻게 (핵심)",1)
P("의도적으로 ML/딥러닝/TSFM을 쓰지 않는다(§4 근거). '모델'은 해석가능한 규칙·선형 합성이다.",b=True)
P("Sleeve ① Equity 시장중립 Long/Short — 다요인 선형 합성",b=True)
BUL([
 "신호: 모멘텀(12-1) + 퀄리티(gp/assets, Novy-Marx) + 실적추정변화 proxy(매출·이익·gp·마진의 전년대비 변화, SF1 datekey PIT)를 각각 횡단면 z-score 후 동일가중 평균.",
 "구성: 섹터중립화 → 상위 20% 롱 / 하위 20% 숏, 동일가중, dollar-neutral. 시장 베타 ≈ 0.",
 "근거: 학습 없는 규칙 기반(과적합 최소). 단순 단일팩터는 v1에서 전멸했고(아래), 결합도 약함 → 이 sleeve는 '수익엔진'이 아니라 '무상관 분산기'(net Sharpe ~0.3).",
])
P("Sleeve ② ETF 추세추종 — 시계열 모멘텀 규칙",b=True)
BUL([
 "신호: 15개 美 ETF(SPY/QQQ/IWM/섹터 SPDR)가 200일 이동평균 위면 롱, 아래면 현금(long-flat). SPY 20일 실현변동성↑이면 노출 ½(vol-regime).",
 "근거: 시계열 추세는 문헌상 강건(Hurst/Ooi/Pedersen·Moskowitz). 단 우리 구현은 주식연계 ETF·200일 룰이라 '강건한 분산선물 추세'의 *취약한 부분집합*. net Sharpe ~0.85지만 2017–26 강추세 레짐 flattered.",
 "이게 책의 주 수익원(질문 2 답: 맞다 — ETF·롱 편향·추세가 수익을 견인. equity sleeve는 분산용).",
])
P("Sleeve ③ LETF 컨벡스 dose — 3배 ETF로 상승 증폭",b=True)
BUL([
 "추세 ON인 지수를 3x ETF(TQQQ/UPRO)로 표현. 실제 LETF 가격(일간리셋 감쇠 내장) 사용 — 3×지수 합성 금지.",
 "고변동이라 위험예산 소량(dollar 비중 ~4%). 결합 시 동일 낙폭에서 연 +1.2% 추가.",
])
P("결합·사이징·리스크(바닥)",b=True)
BUL([
 "결합: 세 sleeve가 거의 무상관 → 위험예산(역변동성) 가중, 추세-복합 ≤65% cap.",
 "사이징: 변동성 목표 + fractional-Kelly(0.5) + 하드 최대레버리지 + 드로다운 ladder(−10/−15/−20/−25%→노출 축소) + cooldown(킬 후 즉시 재가동 금지).",
 "리스크 엔진(결정적): 데이터무결성·드로다운·변동성폭발·시장중립붕괴·모델decay 킬. look-ahead 제거(신호 t→진입 t+1).",
])
P("탈락시킨 것(가짜알파 거부 = 규율 작동)",b=True)
BUL([
 "VRP 단기변동성 매도: 최악월 −62%(Brexit 갭)·MDD −92%, 크래시킬이 단일일 갭 못 막음 → REJECT.",
 "단기 reversal: 비용 전에도 음수·턴오버 937% → REJECT.",
 "ML(ridge/GBM): OOS에서 선형 합성 못 넘음 → 제외(§4).",
])

H("4. 왜 ML/TSFM을 안 썼나 (질문 5)",1)
BUL([
 "TSFM/딥러닝: 사용 안 함. DR-4 모델맵 + council 결정 = alpha 용도로 reject(과적합 위험·해석불가·thin 데이터).",
 "근거 ①(자체 실증): walk-forward OOS bake-off에서 ridge +0.15·GBM +0.02 net Sharpe — 둘 다 단순 선형 합성(0.33) 못 넘음.",
 "근거 ②(deep-research 105에이전트·literature): Gu-Kelly-Xiu(2020) ML 우위는 월 OOS R² 0.33–0.40%로 작고, 헤드라인 Sharpe는 비용 전·마이크로캡 집중(ex-microcap 66–78%↓·deep은 distressed 제외 시 전멸). 즉 ML 엣지는 '솔로가 싸게 거래 못 하는 곳'에 산다.",
 "근거 ③(Codex 적대리뷰): Sharadar 일별 데이터에 deep=theater. → 모델공간 CONSTRAIN(정규화선형+GBM challenger까지·trial ledger·deep 금지).",
 "기술스택: 횡단면 z-합성·200d SMA 추세·vol-targeting·fractional-Kelly·섹터중립·위험예산 결합·walk-forward·deflated-Sharpe/PBO 인지·사전등록 kill·PIT/survivorship·next-bar 실행. (전통적·해석가능)",
])

H("5. 결과 (정직)",1)
P("v1 — 공개 단순 신호 6개 family 전멸(사전등록 kill):",b=True)
TBL([["가설","결과"],["단순 level 팩터","DEAD"],["저DoF 결합","DEAD"],["small/mid 비용후","DEAD"],
     ["event/surprise proxy","FAIL"],["True PEAD(가격반응)","FAIL(반전)"],["analyst revision","데이터 게이트 HOLD"]])
P("v3 — 결합 book(next-bar, 비용후, paper):",b=True)
TBL([["버전","Sharpe","CAGR","MDD","비고"],
     ["결합 무레버","1.18","+8.2%","−10%","무상관 분산"],
     ["공격(vol40·4x)","1.12","+14.3%","−17%","₩100M→3.37억(9년)"],
     ["3개월 시뮬(2026-03~)","(연1.0)","+6.6%","−12%","같은기간 SPY +9.9%(강세장선 책이 못이김)"]])
P("정직: 헤드라인은 backtest·추세 sleeve 견인·강추세 레짐 flattered. forward 정직 기대 Sharpe ~0.5–0.6.",i=True,c=(0x92,0x40,0x0e))

H("6. 3개월 시뮬은 정직했나 (질문 4)",1)
BUL([
 "실행: look-ahead 제거 — 신호는 그 시점까지 데이터로만(PIT), 진입은 next-bar(신호일 다음 거래일). 유니버스는 PIT 멤버십(현재 리스트 소급 금지)·survivorship-safe·배당조정가.",
 "단 정직한 한계: '전략 설계'(어떤 sleeve를 쓸지, 추세가 통한다는 것)는 최근 역사를 *알고* 골랐다 → 파라미터 튜닝은 없지만 *선택*엔 사후정보가 있다. 따라서 3개월 숫자는 '실행상 무누수'지 '순수 blind forward'는 아니다.",
 "진짜 blind 검증 = 오늘부터 forward paper 누적(forward_book.py)뿐. 3개월 +6.6%를 미래 기대값으로 보면 안 됨.",
])

H("7. 이걸 어떻게 믿고 실제 돈을 넣나 (질문 3)",1)
P("정직한 답: 지금은 넣지 않는다 — 시스템이 막도록 설계됐다.",b=True)
BUL([
 "promotion gate(실거래 전 필수): 12개월+ forward 라이브-페이퍼 실적·live Sharpe≥0.8·MDD≤15%·2x비용서 Sharpe≥0.5·데이터무결성 사고 0·용량 충분. 현재 book은 이 게이트를 *의도적으로 통과 못 함*(forward 0개월·Sharpe 0.33 등).",
 "신뢰는 backtest가 아니라 forward 실적이 쌓는다. 경로: forward paper(매월 forward_book) → 게이트 통과 시 사람 더블 사인오프 → 첫 live는 의도자본의 ≤10%.",
 "공격 레버(−18% MDD)는 promotable 한도(MDD≤15%)와 상충 → 승격하려면 저레버로. '공격적'과 '안전 승격'은 트레이드오프.",
])

H("8. 정직한 한계 + 차별점 (질문 5)",1)
BUL([
 "수익은 ETF 추세(롱 편향)가 견인 → 사실상 '주식 베타 타이밍'일 위험. equity 시장중립은 약함(분산용). 강세장에선 SPY 단순보유에 짐.",
 "데이터창 2016–26(한 레짐·2008 미포함). 무상관(+0.06~−0.25)은 benign 샘플일 수 있음(스트레스 시 +0.25~+0.60).",
 "차별점은 '비밀 알파'가 아니라 *규율*: 사전등록 kill·deflated-Sharpe/PBO·정직한 negative·과적합 거부·결정적 risk veto·paper 게이트. 대부분의 리테일 퀀트가 과적합 후 라이브로 가는 것을 이 시스템은 막는다.",
])

H("9. 레퍼런스·실행",1)
P("참고(질문 5): DR-1~4 사전 락 시퀀스 · deep-research(ML vs linear, 트렌드, 검증) · prior-art(Gu-Kelly-Xiu, Avramov, Harvey-Liu-Zhu, Bailey-Lopez de Prado, Hurst/Moskowitz, DeMiguel 1/N, Novy-Marx) · Codex 적대 council.")
P("실행: python engine/data.py validate(무결성) · python engine/_smoke_run.py(동결 재현) · python engine/book_final.py(책) · python engine/forward_book.py(현재 목표+forward log) · python engine/dashboard.py(대시보드).")
P("코드 phase1a/engine/ · 문서 integrated/ · 데이터(라이선스 Sharadar)는 로컬·비공개.",i=True,c=(0x6b,0x72,0x80))

out=os.path.join(data.PHASE1A,"..","integrated","PRAMANA_Technical_Report.docx")
d.save(out); print(f"✅ 기술보고서 → {os.path.abspath(out)} ({os.path.getsize(out)//1024}KB, {len(d.paragraphs)} 단락)")
