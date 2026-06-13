#!/usr/bin/env python3
"""PRAMANA 교수 제출용 최종 보고서 (DOCX) — 8세대 연대기 학술 포장.
negative result + 엄격한 방법론 = 정당한 연구 기여. python engine/build_professor_report.py"""
import os
from docx import Document; from docx.shared import Pt,RGBColor; from docx.enum.text import WD_ALIGN_PARAGRAPH
HERE=os.path.dirname(os.path.abspath(__file__)); REPO=os.path.dirname(os.path.dirname(HERE))
OUT=os.path.join(REPO,"PRAMANA_V4","PRAMANA_Final_Report_for_Submission.docx")
d=Document()
def H(t,l=1,c="1F3864"):
    h=d.add_heading(t,level=l)
    for r in h.runs: r.font.color.rgb=RGBColor.from_string(c)
def P(t,b=False,sz=10.5,it=False,c=None):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=b; r.italic=it; r.font.size=Pt(sz)
    if c: r.font.color.rgb=RGBColor.from_string(c)
    return p
def bullet(t,sz=10):
    p=d.add_paragraph(style="List Bullet"); p.add_run(t).font.size=Pt(sz)
def table(hdr,rows,style="Light Grid Accent 1"):
    t=d.add_table(rows=1,cols=len(hdr)); t.style=style
    for i,h in enumerate(hdr): t.rows[0].cells[i].paragraphs[0].add_run(h).bold=True
    for row in rows:
        c=t.add_row().cells
        for i,v in enumerate(row):
            c[i].text=str(v)
            for pp in c[i].paragraphs:
                for rr in pp.runs: rr.font.size=Pt(9)
    return t

d.add_heading("PRAMANA",0)
P("솔로 + AI 체계적 주식 검증 운영체계(Systematic Equity Validation OS)와",True,13)
P("사전등록 가설 8세대의 정직한 Negative Results에 관한 연구",True,13)
P("— 레버리지·예측 모델 없이 SPY/QQQ 위험조정 초과 알파의 부재를 실증하고, 재사용 가능한 검증 프레임워크와 생존형 코어를 도출하다 —",False,10,it=True,c="808080")
P("PAPER only · NO LIVE · 가상자본 ₩100M · 2026-06-13",False,9,c="808080")

H("초록 (Abstract)",1)
P("본 연구는 개인(solo) 연구자가 AI(LLM)를 보조 엔진으로 활용하여, 미국 주식·ETF를 대상으로 *검증 가능한* 체계적 알파를 탐색·검증하는 운영체계 PRAMANA를 구축한 기록이다. 핵심 방법론은 '데이터 정직성 → 비용 정직성 → 검증 통과 → 그 다음에야 모델 비교'라는 순서 잠금과, 모든 가설에 대한 **사전등록(pre-registration)·OOS·비용후·look-ahead 차단·적대적 AI 검수(adversarial council)·사전 kill 조건**의 일관 적용이다. 단순/선형 횡단면 팩터, 결합·트리·ML, 레버리지 베타, 4종의 마켓타이밍, intraday 급등주 setup, 그리고 정성(8-K 공시) 텍스트에 이르기까지 **8세대의 가설을 유료 기관급 PIT 데이터(Sharadar Core US Equities·survivorship-free)와 무료 보조(EDGAR 8-K·yfinance forward)로 검증한 결과, **이 데이터·비용·제약·기간(2016–2026)에서 개인이 SPY/QQQ를 위험조정 기준 반복적으로 초과하는 '매수형 알파'는 발견되지 않았다(범위 조건부 negative — 보편적 부재 주장이 아님).** (저품질 데이터 탓이 아니라 — 유료 기관급 데이터로도 부재한다는 점이 negative를 더 강하게 만든다.) 이는 efficient market 가설, 최근 SPIVA U.S. scorecards(대형주 액티브 펀드 다수가 인덱스 미달), Gu-Kelly-Xiu(2020 — 전문 대규모 ML에도 월 OOS R²가 ≈0.33–0.40%로 작고 구현 장벽이 큼)와 정합하는 *scope-conditional negative*(8개 사전등록 시도 범위 — 보편적 부재 주장이 아님)이다. 그러나 본 연구는 (1) 재사용 가능한 검증 OS, (2) 자체구축 PIT 벤치마크(실제 SPY와 corr 0.998), (3) 위험조정상 우월한 생존형 분산 코어(V7 4-sleeve), (4) 일관된 '악재 공시 회피' 필터, (5) 가짜 알파를 자본 투입 전에 제거하는 규율을 산출하였다.",False,10.5)

H("1. 연구 배경 및 목적",1)
P("'좋은 종목을 찍는 AI'가 아니라, **개인/소규모 팀이 AI와 함께 운영 가능한 체계적 주식 연구·검증·페이퍼 트레이딩 OS**를 구축하는 것이 목적이다. 통상의 퀀트 연구가 '모델을 먼저 고르고 백테스트가 좋게 나오는' 순서로 인해 비용 과소추정·미래정보 누수·생존편향에 취약하다는 문제의식에서 출발하여, 검증을 *모델보다 앞에* 두는 역순 설계를 채택했다.",False,10.5)

H("2. 방법론 (Methodology) — 본 연구의 핵심 기여",1)
P("알파의 *발견*이 아니라 알파 주장의 *반증 가능성(falsifiability)*을 방법론의 중심에 두었다.",False,10.5,it=True)
for t in ["순서 잠금: data honesty → cost honesty → validation → then models (모델은 마지막).",
          "사전등록 + trial ledger: 핵심 kill-test는 결과를 보기 *전에* 가설·비교군·kill 조건을 문서로 고정했고, 나머지 반복은 trial ledger + 사후 변경금지 규율로 추적(goalpost 이동 = config-mining 금지).",
          "Look-ahead 차단: next-bar 체결, PIT(point-in-time) 멤버십, acceptance-timestamp 기준 진입.",
          "OOS·비용후: walk-forward + 최종 untouched holdout, 거래비용·턴오버·세금 반영.",
          "적대적 AI 검수(Adversarial Council): Claude(설계) ↔ Codex(반증) 2모델이 look-ahead·survivorship·overfit·hidden-coupling을 상호 공격(no-echo).",
          "다중검정 위험 통제: trial registry + 사전 kill 조건으로 1차 완화 (DSR/PBO는 정식 산출 미완료·향후 과제 — 개념적 인지 수준).",
          "판정 라벨 분리: capital-veto(자본권한)와 research-veto(연구개방)를 분리 — 연구는 자본권한 0으로 개방."]:
    bullet(t)

P("2.1 검증 OS의 핵심 증거 — '가짜 알파' 2건 적발·폐기",True,10.5)
P("이 OS의 가치는 통과시킨 것이 아니라 *잡아낸 가짜 알파*에 있다. 적대적 검수(Codex)가 두 번의 look-ahead 누수를 적발해 '좋아 보이던' 후보를 폐기했다 — 검증 장치가 실제로 작동함을 보이는 핵심 증거다.",False,10.5,it=True)
for t in ["Receipt 1 (RVOL look-ahead · intraday): 초기 sim이 당일 전체 거래량을 진입 전에 사용(미래정보). 진입시점 누적 RVOL로 수정 후 재실행 → 엣지 붕괴(중앙값 −0.41%·false breakout 56%) → setup 폐기(DEAD).",
          "Receipt 2 (동적 allocator same-day leak): 동적 배분이 당일 신호로 당일 수익을 계산. next-bar 수정 + ablation → 동적 기여 −113%p → 동적 타이밍 OFF 확정.",
          "함의: 두 건 모두 누수 제거 후 '좋던' 결과가 사라졌다 — backtest 수익률보다 leakage/overfitting 방어가 본 연구의 실질이다."]:
    bullet(t)

H("3. 실험 연대기 (8세대) — 무엇을 검증했고 무엇이 죽었나",1)
P("[네이밍 안내] 'PRAMANA'는 프로젝트명이고, v1~v8은 *전략 북(book)의 진화 세대*이다(프로젝트 버전이 아님). v1~v3은 횡단면·풀북 단계, v4~v8은 core-satellite 재설계 이후의 북 이터레이션(v4 Core Beta → v7 4-sleeve), A1은 최종 공격형 북이다. 저장소 폴더명 'PRAMANA_V4'는 core-satellite 재설계 시점의 *작업공간 라벨*일 뿐, 최신 북 버전(v7)이나 프로젝트 버전과 무관하다.",False,9.5,it=True,c="808080")
P("아래 8세대는 5개 의사결정 체크포인트로 압축된다 — 각 체크포인트는 직전 실패가 강제한 다음 가설이며, 알파 탐색이 좁아지는 깔때기다.",False,10.5)
table(["체크포인트","가설","결과 → 다음 가설을 강제"],[
 ["CP1 횡단면 팩터군","value/mom/quality/lowvol + ML","IC≈0(쉬운 알파 없음) → 레버리지로"],
 ["CP2 풀 레버리지 북","trend+LETF·VRP·레버드 베타","SPY 위험조정 초과 실패(레버 ≠ 알파) → 구조 분산으로"],
 ["CP3 코어-새틀라이트(4-sleeve)","구조적 분산 코어","베타지 알파 아님(위험효율) → 타이밍으로"],
 ["CP4 마켓타이밍 ×4","regime/throttle/derisk/ladder","4전 4패(후행신호 벽) → 정성·이벤트로"],
 ["CP5 A1/A2 공격 북","catalyst·비대칭 베팅","검증 알파 없음 — 정직한 위험베팅(에필로그)"],
])
P("(세부 원장 — 8세대 전체 판정)",False,9,it=True,c="808080")
table(["북 세대","가설/접근","판정","핵심 근거"],[
 ["v1","단순 횡단면 팩터(value/momentum/quality/lowvol)","FAIL","Rank IC ≈ 0 · quality decay(0.22→0.046)"],
 ["v1","결합·ridge·GBM·tree","FAIL","OOS net vs cap-weight 음수 (GKX와 정합)"],
 ["v1","event/earnings drift(숫자 surprise)","FAIL","net vs cap-weight −0.90%"],
 ["v3","trend+LETF 위성·VRP·mean-reversion","REJECT/노이즈","+0.15%/yr 노이즈 · tail −92% · turnover 3660%"],
 ["v4–v6","Core Beta · 레버드 베타 · 분산","알파 아님","v5 Sharpe≈QQQ(레버지 알파 아님)"],
 ["v7","4-sleeve 분산 코어","생존코어(채택)","Sharpe 1.21·MDD −18% / 단 누적 절반"],
 ["v8","Levered 4-sleeve","REJECT","닷컴 proxy −49%(레버 꼬리)"],
 ["—","마켓타이밍 4종(regime/throttle/derisk/MT-1)","4전 4패","모두 static 4-sleeve 미달(후행신호 벽)"],
 ["Alpha","intraday 급등주 setup(ORB/VWAP/RVOL)","DEAD","RVOL look-ahead 누수·강세장 베타"],
 ["Alpha","8-K 정성 catalyst(POS) / (NEG)","POS FAIL / NEG 일관","사는 알파 없음 / 악재 회피는 일관 신호"],
])

H("4. 핵심 결과 (Results)",1)
P("[성과 상태 정의 — 정직성] 아래 수치의 성격을 명확히 구분한다(backtest를 live 실적으로 오인 방지).",False,9.5,it=True,c="808080")
table(["상태","정의","V7 현황"],[
 ["Backtest","과거 데이터 in-sample 재현(미래정보 차단)","2019~·비용후 — 아래 §4.1 수치 전부 이것"],
 ["Warmup","forward 시작 전 캐시 적재","600d yfinance(대시보드 × 곡선 맥락·₩NAV 아님)"],
 ["Forward paper","인셉션 이후 무인 관찰(자본 0)","인셉션 2026-06-11·실질 1~2거래일 — 성과 판단 근거 미달"],
 ["Live capital","실제 자본 투입","0 (사람 게이트 전 영구 금지)"],
])
P("→ §4.1은 전부 Backtest다. forward paper는 1~2거래일로 성과 판단 근거가 아니며, live capital은 0이다. 무인 cron은 등록됐으나 실제 트리거는 미검증(live-ready, not live-validated).",False,9.5,it=True,c="808080")
P("[Operation Receipt] 작동성은 산출물 존재가 아니라 실행으로 증명한다 — 실제 명령·산출·한계:",False,9.5,it=True,c="808080")
table(["구성요소","실행 명령","마지막 산출·상태","한계"],[
 ["검증 파이프 B0","bash reproduce.sh (무료·키 불필요)","B0 machinery PASS · self-built PIT vs 실제 SPY corr 0.998","free=survivorship-biased smoke"],
 ["V7 paper forward","engine/forward_runner_v7.py","forward 1~2거래일 · 자본0 · v7_forward_dashboard.html","cron 트리거 미검증(수동 실행)"],
 ["A1/A2 paper book","engine/a1_live_runner.py · a2_live_runner.py","health.json ok · append-only NAV 로그","positions JSON 비면 cash"],
 ["재현성 동결","b0 frozen snapshot (data_hash)","동일 해시 재현 PASS","paid 결과 재현은 Sharadar 키 필요"],
])
P("4.1 멀티앵커 성과 (비용후·in-sample/backtest) — 분산 코어 V7 vs 인덱스",True,10.5)
table(["진입 시점","V7 4-sleeve","QQQ","SPY"],[
 ["3개월","+4.6% / −5% / Sh 1.41","+20.2% / −7% / 3.42","+11.1% / −6% / 2.84"],
 ["6개월","+8.4% / −7% / 1.29","+14.9% / −12% / 1.54","+7.7% / −9% / 1.17"],
 ["12개월","+26.5% / −7% / 2.13","+35.3% / −12% / 1.85","+24.1% / −9% / 1.82"],
 ["풀(2019~)","+174.9% / −18% / 1.21","+305% / −35% / 0.94","+186% / −34% / 0.86"],
])
P("→ V7은 모든 구간에서 누적수익을 인덱스에 양보하는 대신 MDD를 절반으로 줄이고 위험조정(Sharpe)을 개선한다. 즉 '초과수익(알파)'이 아니라 '위험효율(분산 프리미엄)'이며, 최대복리 목적함수에서는 인덱스가 우월하다.",False,10,it=True)
P("4.2 메타 패턴 (가장 중요한 발견)",True,10.5)
P("모든 횡단면 신호군이 1/N(동일가중)은 이기지만 cap-weight(시가총액 가중) 인덱스는 이기지 못했다. 일관된 원인은 2016–2026년의 '메가캡 지배 레짐'이다. 횡단면 틸트는 *신호의 우열과 무관하게* 이미 지배적인 소수 초대형주를 덜 담게 되므로 cap-weight 인덱스에 구조적으로 진다. 즉 패배는 '신호 부재'가 아니라 '벤치마크 구조'에서 기인한다.",False,10.5)
P("4.3 정성(8-K) 결과 — 비대칭",True,10.5)
P("좋은 공시(계약·실적) 매수는 OOS에서 실패(이미 가격 반영)한 반면, 나쁜 공시(희석·회계 문제·상폐위험)는 후폭풍이 수일간 지속되어 −0.75%의 일관된 음(−)의 drift를 보였다. → '사는 신호'는 없고 '피하는 필터'는 존재한다.",False,10.5)

H("5. 논의 (Discussion) — 왜 개인이 어려운가",1)
for t in ["정보이론적 한계: 모델은 정보를 *추출*할 뿐 *창조*하지 않는다. 기성·정형 데이터(유료 PIT 포함)에 잔존 알파가 작으면 어떤 ML도 한계를 넘지 못한다(GKX 월 OOS R² ≈ 0.4%).",
          "벤치마크의 강도: SPY/QQQ는 '평균 개인'이 아니라 시장 집단지성·초저비용·세금효율·승자 비중확대를 내장. 프로 79%(SPIVA)도 미달.",
          "산업 vs 개인: XTX·Jane Street 등은 단일 모델이 아니라 미세구조·유동성공급·초저지연 체결·독점데이터·자본·조직의 *시스템*으로 수익을 낸다. 개인이 접근 가능한 데이터(유료 기성 PIT 포함)로는 복제 불가능한 게임이다.",
          "개인의 잔존 우위: (1) 행동 우위(폭락에서 매도하지 않는 규율), (2) 제약 우위(기관이 못 드는 극소형), (3) 정성 해석(LLM 보조), (4) 리스크 구조(생존형 분산). — 단 (1)·(4)가 실현수익에 가장 크게 기여."]:
    bullet(t)

H("6. 결론 (Conclusion)",1)
P("유료 기관급 PIT 데이터(Sharadar)로도 개인이 SPY/QQQ를 위험조정 기준 반복 초과하는 매수형 알파는 발견되지 않았다(8개 사전등록 시도의 *scope-conditional negative* · 저품질 데이터 탓 아님). 본 연구는 실패가 아니라 *현재 데이터·비용·기간·제약(2016–2026)에서 반복 가능한 매수형 알파의 부재를 엄격히 실증*한 범위 조건부 결과이며, 산출물은 ① 재사용 검증 OS, ② PIT 벤치마크(corr 0.998), ③ backtest 생존코어 후보 V7, ④ 악재 공시 회피 필터, ⑤ 가짜 알파를 자본 투입 전 제거하는 규율이다. (에필로그) 후속 페이퍼 북 A1(Catalyst Attack)·A2(Convex Raider)는 위험을 정직하게 인정한 paper·자본0 공격 베팅으로, *검증된 알파가 아니며 본 연구의 기여 주장도 아니다(not a contribution claim)*.",False,10.5)

H("7. 한계 및 향후 과제 (Limitations)",1)
for t in ["표본 기간(2016–2026)은 메가캡·저변동성 레짐에 편향 — 닷컴/2008은 proxy로만, forward MDD 과소추정 가능.",
          "in-sample/benign sample 위의 결론 — 모든 채택은 paper·자본권한 0·12개월 forward·2-feed reconciliation·사람 게이트 전 자본 금지.",
          "정성(LLM) 축은 baseline까지만 검증 — delayed recognition·LLM catalyst grade는 forward 관찰 대상.",
          "공매도·옵션·극소형은 솔로 실행 제약으로 미검증."]:
    bullet(t)

H("8. 학술적 기여 (Contribution)",1)
P("본 연구의 기여는 '알파를 찾았다'가 아니라 '솔로가 유료 기관급 데이터(Sharadar)까지 동원한 환경에서 알파의 부재를 *사전등록·OOS·비용후·적대적 검수로 (trial registry·kill-gate 수준에서) 강하게 실증*하고, 재사용 가능한 검증 프레임워크를 구축했다'는 데 있다. 특히 적대적 검수가 look-ahead 누수를 2회 적발해 가짜 알파를 폐기한 것(§2.1)은 본 OS가 '통과'가 아니라 '기각'으로 작동하는 false-alpha rejection machine임을 보인다. Negative result이되 방법론적 엄밀성과 재현성(clone 후 one-command 재현·데이터 해시 동결)을 갖춘 정직한 연구이며, AI(LLM)를 *예측 엔진이 아니라 적대적 검수·정성 구조화 보조*로 위치시킨 운영 모델을 제시한다.",False,10.5)

H("부록 A. 다중검정 통제 요약 (Trial Registry · DSR/PBO)",1)
P("선택편향·overfit 통제 현황. 정식 수치 산출이 아닌 항목은 'conceptual(개념 적용)'로 명시해 과대주장을 피한다.",False,9.5,it=True,c="808080")
table(["통제 항목","적용","비고"],[
 ["테스트한 전략 family","8세대(팩터·결합/ML·레버·분산·타이밍×4·intraday·8-K)","family별 사전등록 protocol 기록"],
 ["사전등록 kill 조건","결과 보기 전 잠금(goalpost 고정)","Broad_Universe·AlphaLab·Crashpack protocol"],
 ["next-bar·PIT·OOS·비용후","전 실험 일관 적용","look-ahead 2건 적발·수정(§2.1)"],
 ["Deflated Sharpe / PBO","conceptual (개념 적용·전수 산출표는 미작성)","다중검정 인지 · 정식 계산은 향후 과제"],
])

P("",False,6)
P("부속 자료: PRAMANA_Lineage_Dossier.docx(v1~V8 계보·LOCK) · PRAMANA_All_Experiments_Ledger.md · PRAMANA_Conclusions_OneLine.md · PRAMANA_A1_Attack_Book_Final.md · 라이브 대시보드(outputs/*.html) · 코드/데이터 repo(PAPER, no live).",False,8.5,it=True,c="808080")
H("부록 B. Evidence Ledger — 사전등록 8세대 trial 추적",1)
P("각 시도의 사전등록 문서·데이터(PIT)·IS/OOS·비용·벤치·판정·kill·산출물 경로. negative result의 재현·감사를 위한 단일 인덱스(이 한 장이 '주장'을 '기록'으로 만든다).",False,9.5,it=True,c="808080")
table(["Trial","사전등록·데이터(PIT)","IS/OOS·비용","벤치","판정 / kill","산출물 경로"],[
 ["v1 팩터","B2~B5 protocol · Sharadar PIT","walk-fwd · 비용후","cap-weight","FAIL (Rank IC≈0)","reports/B2B5_broad_result.md"],
 ["v1 ML","결합·ridge·GBM·tree","OOS holdout · 비용후","cap-weight","FAIL (OOS net<0)","reports/Phase1B_lowdof_result.md"],
 ["v1 event","8-K surprise · EDGAR","next-bar","cap-weight","FAIL (−0.90%)","reports/US_event_drift_result.md"],
 ["v3 풀북","trend/LETF/VRP/MR","비용후 · turnover","SPY","REJECT (노이즈·tail)","reports/V3_sleeve_vrp_result.md · V3_overlay_alpha_result.md"],
 ["v4~v6","Core/레버/분산 · yfinance","멀티앵커 · 비용후","SPY/QQQ","알파 아님 (Sharpe≈QQQ)","PRAMANA_V5_Problem_Frame_v0.2.md"],
 ["v7 4-sleeve","test_4sleeve · 2019~","비용후","QQQ","backtest 생존 후보 (paper)","reports/MT1_result.md · V8_Candidate_crashpack_result.md"],
 ["v8 레버","crash-pack proxy","2000/2008 stress","A1 MDD cap","REJECT (닷컴 −49%)","reports/V8_Candidate_crashpack_result.md"],
 ["Alpha Lab","AlphaLab v1/v2 protocol","entry-time · forward","SPY","DEAD / forward 관찰","reports/AlphaLab_v1_result.md"],
])
P("→ 핵심 kill-test는 결과 전 사전등록(kill 조건 잠금)됐고, 모든 trial의 판정·산출물 경로가 repo에 보존된다. 사후 변경금지 규율이 'goalpost 이동(config-mining)'을 막는다.",False,9.5,it=True,c="808080")

H("References",1)
for t in ["SPIVA U.S. Scorecard — S&P Dow Jones Indices (액티브 펀드 다수가 인덱스 미달, ~79%).",
          "Gu, Kelly & Xiu (2020), Empirical Asset Pricing via Machine Learning, Review of Financial Studies (ML 월 OOS R² ≈ 0.33–0.40%).",
          "Bailey & Lopez de Prado (2014), The Deflated Sharpe Ratio, Journal of Portfolio Management (다중검정·선택편향 보정).",
          "Harvey, Liu & Zhu (2016), ...and the Cross-Section of Expected Returns, Review of Financial Studies (false discoveries · 임계 t).",
          "Lopez de Prado (2018), Advances in Financial Machine Learning, Wiley (PBO · purged/embargoed CV)."]:
    bullet(t,9)
d.save(OUT)
print("✅ 교수 제출용 보고서:",OUT)
if __name__=="__main__": pass
