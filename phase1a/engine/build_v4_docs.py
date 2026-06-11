#!/usr/bin/env python3
"""PRAMANA V2 전환 산출물 3종 DOCX 생성:
1) V1 진단 사후분석(설계·전략·모델·근거 + 잘못된 부분 빨강 주석)
2) V2 설계도(8레이어·간략 + 합의결정 + 회의안건)
3) 마스터 다이제스트(모든 연구 한 파일 — 연구아크·실증결과·아키텍처·인덱스)
근거: 서브에이전트 3종 디제스트 + 멀티시점 시뮬 + Codex 리뷰 + DR charter 원본."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
ROOT="/home/click/ghq/github.com/click6067-ship-it/quant-pramana"
OUT=os.path.join(ROOT,"PRAMANA_V4"); os.makedirs(OUT,exist_ok=True)
TEAL=(0x0e,0x74,0x90); RED=(0xc0,0x00,0x00); GRAY=(0x6b,0x72,0x80); AMBER=(0x92,0x40,0x0e)

def newdoc():
    d=Document(); f=d.styles["Normal"].font; f.name="Calibri"; f.size=Pt(10.5); return d
def H(d,t,l=1,c=TEAL):
    h=d.add_heading(t,level=l)
    for r in h.runs: r.font.color.rgb=RGBColor(*c)
    return h
def P(d,t,b=False,i=False,c=None,sz=10.5):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=b; r.italic=i; r.font.size=Pt(sz)
    if c: r.font.color.rgb=RGBColor(*c)
    return p
def RED_(d,t):  # ⚠ 사후진단 빨강 주석
    p=d.add_paragraph(); r=p.add_run("⚠ 사후진단: "+t); r.italic=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor(*RED); return p
def BUL(d,items):
    for it in items:
        if isinstance(it,tuple): d.add_paragraph(it[1],style="List Bullet 2")
        else: d.add_paragraph(it,style="List Bullet")
def TBL(d,rows,hdr=True,style="Light Grid Accent 1"):
    t=d.add_table(rows=len(rows),cols=len(rows[0])); t.style=style
    for i,row in enumerate(rows):
        for j,v in enumerate(row):
            c=t.cell(i,j); c.text=str(v)
            for para in c.paragraphs:
                for r in para.runs: r.font.size=Pt(9)
            if i==0 and hdr:
                for para in c.paragraphs:
                    for r in para.runs: r.bold=True
    return t

# ════════════════════════════════════════════════════════════════════
# DOC 1 — V1 진단 사후분석 보고서
# ════════════════════════════════════════════════════════════════════
d=newdoc()
d.add_heading("PRAMANA 사후분석 보고서 (v1–v3)",0)
P(d,"Diagnostic Post-Mortem · 2026-06-11 · PAPER only",i=True,c=GRAY)
P(d,"범례 — 검정: 사실(설계·전략·모델·선택 근거).  빨강(⚠ 사후진단): 지금 시점에서 잘못됐다고 판단되는 부분 + 교정.",b=True)
P(d,"명명: '지금까지의 PRAMANA'(내부 v1·v2·v3 = 연구·검증 엔진 + Phase1A 실증 + 페이퍼 북) 전체 진단. 다음 재설계 = V4(core-satellite).",i=True,c=GRAY)

H(d,"0. 한 줄 결론")
P(d,"V1은 '실패한 코드'가 아니라 '맞게 만들었으나 — 틀린 잣대로 — 틀린 것을 검증한' 시스템이다. 기계(데이터·검증 파이프라인)는 정확히 작동했고(PIT cap-weight가 실제 SPY와 corr 0.998), 단순/선형 팩터엔 진짜 엣지가 없다는 걸 정직하게 입증했다. 핵심 결함은 모델이 아니라 ① 목표/잣대 미스칼리브레이션(기관용 bar를 소액 공격 계좌에 적용), ② 이진 폐기의 경직성, ③ core 없이 위성만 섞어 SPY를 못 이긴 구조, ④ 코드 전 문서 과중(analysis-paralysis)이다.")

H(d,"1. 설계 구조 (8 레이어)")
TBL(d,[["레이어","구현 모듈","역할"],
 ["1 데이터","data.py","Sharadar SEP/DAILY/SF1/SFP/TICKERS 캐시·manifest/hash·PIT·survivorship·CA/배당"],
 ["2 신호/모델","features.py","동결 6팩터(value·quality·momentum·lowvol·event)+추세/vol, walk-forward 프레임"],
 ["3 알파 앙상블","combined_book.py","메타신호(z-mean)→무상관 sleeve 결합 (핵심 레버)"],
 ["4 포트폴리오","combined_book.py","quintile-EW / 50-50 (최적화 DEFER — Codex: 효과<리스크)"],
 ["5 리스크엔진","risk_engine.py","결정적 veto·fractional-Kelly·vol-target·4 kill·gap추정"],
 ["6 실행","execution.py / simulator.py","per-name 체결(스프레드/수수료/차입/충격/부분체결)"],
 ["7 모니터/kill","monitor.py","일일지표 A-O·K1-6 graded kill·heartbeat·promotion gate"],
 ["8 거버넌스","book_runner.py","paper→live 사람 게이트·LLM off-path·29단계 일일 spine"]])
P(d,"엔진 spine 실행: 110 리밸런스·CAGR +3.24%·Sharpe 0.36·maxDD −18.1%·평균레버 1.00x·kill 50회(corr_breakdown 26·model_decay 24). 전부 config-driven·API-free·동결 재현.",i=True,c=GRAY)
RED_(d,"8 레이어 자체는 건전. 문제는 이 구조가 'SPY/QQQ 코어를 충분히 들고 가는 공격형'이 아니라 '시장중립/추세/LETF를 위험예산으로 섞는 연구형'이라는 것 — core beta가 빠져 강세장에서 SPY에 구조적으로 진다. V4는 core-satellite로 교정.")

H(d,"2. 사용한 전략·모델 + 근거")
P(d,"2.1 동결 6팩터 (의도적 boring baseline — 채택용이 아니라 죽이기용)",b=True)
TBL(d,[["팩터","정의","근거"],
 ["value","1/pb","고전 가치, 횡단면 baseline"],
 ["quality","gp/assets (PIT by datekey)","Novy-Marx 수익성"],
 ["momentum","12-1","Jegadeesh-Titman"],
 ["lowvol","−126d 변동성","low-vol anomaly"],
 ["event","매출·이익·gp·마진 YoY z-평균","surprise proxy"],
 ["(baseline)","1/N 동일가중","DeMiguel 2009 = 진짜 boring baseline(추정 MVO 아님)"]])
P(d,"결과: 전부 DEAD (사전등록 비용후 kill). momentum net +7.6%지만 IC-IR 0.10·turnover 301%; quality만 broad에서 SURVIVE-screen(IC-IR 0.220) 후 quarantine FAIL.",i=True,c=GRAY)
RED_(d,"단순팩터 사망은 '잣대가 엄해서'가 아니라 IC≈0·비용후 음수 = 진짜 엣지 없음. 여기엔 교정 불필요(정직한 negative). 단 quality는 long-only가 cap-weight 못 이긴다는 *기관식 잣대*로 FAIL 줬다 — 소액·절대수익 관점이면 '코어 대비 한계 무상관 기여' 잣대로 재심사할 가치 있음(V4 위성 후보). 단 최근 절반 decay(IC-IR 0.42→0.046)는 실재 우려.")

P(d,"2.2 v3 3 sleeve (페이퍼 북)",b=True)
TBL(d,[["sleeve","모델","결과","역할"],
 ["equity MN L/S","mom+qual+rev 선형합성·섹터중립·롱숏","Sharpe 0.31·CAGR +2.6%·MDD −21%","분산용(약함)"],
 ["ETF trend","200d-SMA long-flat·15 ETF+vol gate","Sharpe 0.85·CAGR +9.7%·MDD −16%","수익엔진"],
 ["LETF","실제 3x ETF(TQQQ/UPRO) 경로","combo 매칭DD서 +1.21%/yr","상승 dose"]])
P(d,"결합: 무상관(corr +0.06)→위험예산 가중→레버. 근거(council LOCK): 진짜 레버는 모델 복잡도가 아니라 분산+실행/리스크 엔지니어링.",i=True,c=GRAY)
RED_(d,"equity sleeve(Sharpe 0.31, 46% 비중)가 trend 엔진(0.85)을 *희석*했다. 멀티시점 net 재측정 결과 풀북은 3M/6M/1Y/풀사이클 어느 구간도 SPY를 위험조정으로 못 이긴다(풀사이클 3x +699% vs SPY +307%는 레버일뿐 — MDD −45% vs −34%, Sharpe 0.82<0.90). trend+LETF만(드래그 제거) 돌리면 ~400일 +22% vs SPY +15%·Sharpe 1.38 → 엔진은 trend, equity는 빼야 했다.")

P(d,"2.3 탈락시킨 것 (가짜알파 거부 = 규율 작동)",b=True)
BUL(d,["VRP 단기변동성 매도: 최악월 −62%·MDD −92% → REJECT(파산 kill).",
 "단기 reversal: gross Sharpe −0.52(비용 전에도 음수)·turnover 937% → REJECT.",
 "True PEAD: 부호 역전(IC −0.0153, 대형주 과반응→반전) → REJECT.",
 "ML(ridge/GBM): OOS서 3-feature 선형(0.33) 못 넘음 → 제외."])
P(d,"NO ML/TSFM 근거(3원 수렴): Codex(일별 Sharadar에 deep=theater)·자체 bake-off(ridge +0.15/GBM +0.02 못 넘음)·deep-research 105에이전트(GKX ML OOS R² 0.33-0.40%/월·비용전·마이크로캡). → 'ML 엣지는 솔로가 싸게 못 거래하는 곳에 산다.'",i=True,c=GRAY)
RED_(d,"NO-ML 결정은 옳다(교정 불필요). 단 DL을 *완전* 배제한 건 과했음 — alpha 예측엔 reject지만 스캐너/필터/메타라벨러(이 신호가 지금 레짐서 먹힐 확률)로는 정당. V4서 보조엔진으로 재배치.")

H(d,"3. 검증 결과 — 무엇이 죽고 살았나")
TBL(d,[["실험","결과","판정"],
 ["B0 broad (S&P500 PIT cap-weight)","실제 SPY와 corr 0.998·+0.55%p/yr","PASS(기계검증)"],
 ["B2/B4/B5 단순팩터","|IC|<0.01","DEAD"],
 ["B3 quality broad → quarantine","IC-IR 0.220 → decay·long-only<CW −1.15%","SURVIVE-screen→FAIL"],
 ["Phase1B 결합(blend/rank/ridge)","ridge OOS IC-IR 0.014·전부 <CW","FAIL(family TERMINATE)"],
 ["event/surprise (대형·small/mid)","no-decay지만 cap-weight 벽+비용","FAIL"],
 ["V3 trend overlay","Sharpe 0.85·corr +0.018","SURVIVE"],
 ["V3 LETF","combo +1.21%/yr(매칭DD)","SURVIVE"],
 ["VRP·reversal","파산·음수","REJECT"]])
P(d,"메타패턴(일관된 킬러): 모든 횡단면 신호가 1/N은 이기지만 cap-weight long-only엔 진다 = 2016-26 메가캡 지배 레짐. 신호는 small/mid에 더 살지만 유동성집중+비용이 거래가능 슬라이스를 죽인다.",i=True,c=GRAY)
RED_(d,"L/S 스프레드는 양수였는데 'long-only가 cap-weight 못 이김' 잣대로 깎인 사례 다수 = 잣대-신호 미스매치. 소액 공격 계좌엔 long-short 절대수익이 유효할 수 있음(단 숏 비용·차입 현실 반영 필수).")

H(d,"4. 핵심 결함 진단 (전부 ⚠ 사후진단)")
RED_(d,"① 잣대 미스칼리브레이션 — promotion gate(Sharpe≥0.8·MDD≤15·cap-weight 이기기)가 기관/수탁자용. 소액 공격 계좌의 목적함수(절대수익·큰 변동 감내)와 불일치 → 살아날 전략을 죽임.")
RED_(d,"② 이진 폐기의 경직성 — 통과/폐기뿐. 'standalone 엣지는 없지만 앙상블/특정 레짐선 쓸모'를 못 살림. V4: live/paper/변형/sandbox/영구폐기 + capital-veto vs research-veto 분리.")
RED_(d,"③ core 부재 — SPY/QQQ 코어 없이 위성만 섞어 강세장에서 SPY에 구조적으로 짐. core-satellite였어야.")
RED_(d,"④ analysis-paralysis — 문서 15개·게이트 14·registry 60필드인데 한동안 코드/데이터 0. 'wrong scale이 모델 아닌 프로세스에 발생'. V4: 크루드한 것부터 ship.")
RED_(d,"⑤ 낙관 보고 — 대시보드/보고서가 'Sharpe 1.1+'를 헤드라인했으나 gross·유리한 구간. net·멀티시점은 SPY 못 이김. 보고는 항상 net·worst-horizon 병기.")
RED_(d,"⑥ Codex 지적(REVISE) — same-close 체결 누수(→next-bar 수정), '인-샘플 equity 베타 타이밍을 분산북으로 포장', 무상관 +0.06은 benign 샘플(스트레스 시 +0.25~0.60), 무료 단일피드=레버 실자본 부적합. forward 정직 Sharpe 0.35-0.55.")

H(d,"5. 유지할 가치가 있는 것 (이전 작업의 자산)")
BUL(d,["데이터 정직성 기계: PIT·survivorship·CA·배당 → cap-weight corr 0.998 (강력 검증).",
 "비용 정직성·사전등록 kill·look-ahead 제거(next-bar)·동결 재현.",
 "NO-ML 결정(council 3원 수렴) — 옳음, 유지.",
 "정직한 negative 문화: 가짜알파를 시스템이 안 만든다(가장 큰 자산)."])
H(d,"6. V4로 넘기는 교훈")
P(d,"검증을 버리는 게 아니라 *잣대를 갈아끼운다*. core-satellite로 SPY와의 싸움을 끝낸다. 위성은 코어 대비 한계 기여(비용후)로 심사. 분봉은 최강 격리. 빅뱅 대신 Core+trend+attribution부터.")
d.save(os.path.join(OUT,"PRAMANA_PostMortem_v1to3.docx"))
print("✅ 1/3 PRAMANA_PostMortem_v1to3.docx")

# ════════════════════════════════════════════════════════════════════
# DOC 2 — V2 설계도
# ════════════════════════════════════════════════════════════════════
d=newdoc()
d.add_heading("PRAMANA V4 — 설계도",0)
P(d,"Core-Satellite + Meta-Allocator · 2026-06-11 · PAPER first",i=True,c=GRAY)
P(d,"목표: SPY/QQQ 코어 + 검증된 공격 위성으로 초과수익 — 사전약속 낙폭 한도 내 생존. 'SPY를 버리고 이기기'(짐 증명) 아님, 'SPY를 깔고 위에 얹기'.",b=True)

H(d,"A. 합의된 핵심 결정 (Claude+GPT 취합·정제)")
BUL(d,["Core-satellite: SPY/QQQ/현금 코어를 기본 장착, 위에 공격 위성으로 초과수익. SPY와의 싸움 종료.",
 "전략군 고정·비중만 레짐 따라 자동 변경(즉흥 변경 금지='방금 지나간 차트에 춤추기' 차단).",
 "리스크엔진: 없애지 않고 *판정 언어* 교체 — live/paper/변형/sandbox/영구폐기 + capital-veto vs research-veto 분리.",
 "DL/TSFM: alpha 왕 아님 → 스캐너/필터/메타라벨러/레짐판별 보조. from-scratch 금지(FinBERT 등 기성).",
 "기관 포지션(13F/버핏): entry 신호 아님(45일 지연) → 느린 코어 틸트 feature. 단일종목 확신≠레짐.",
 "분봉: 더 고도화된 엔진 아니라 *가장 강하게 격리할 실험실*. capital allocator 침범 금지.",
 "위성 심사 = 단독 Sharpe 아니라 *코어 대비 한계 초과수익(비용후)*. 안 그러면 SPY 베타를 비싸게 재현.",
 "목적/제약 순서: 생존(MDD≤한도)=제약, 초과수익=목적. max 초과수익 s.t. MDD≤X.",
 "attribution 필수: core/위성한계/얼로케이터타이밍/분봉 gross-비용 분해 → 자기기만 방지.",
 "빅뱅 금지: Core+trend+attribution+멍청한 allocator 먼저 → 위성 하나씩 → 분봉 lab → DL."])

H(d,"B. 8 레이어 설계 (간략 — 시스템 / 작동원리 / 구조)")
def layer(n,title,system,how,struct):
    P(d,f"L{n}. {title}",b=True,c=TEAL)
    P(d,f"· 시스템: {system}")
    P(d,f"· 작동원리: {how}")
    P(d,f"· 구조: {struct}")
layer(1,"Data & Integrity","일봉(EOD 다소스)+분봉(intraday, lab 전용). PIT·survivorship·CA.",
 "장마감 후 pull→fail-closed 무결성 검사(stale/NaN/바개수/배드틱)→통과시만 갱신.","실자본=2 독립피드+브로커 대조. 페이퍼=무료EOD 허용+health.json. append-only 동결.")
layer(2,"Regime Engine","일봉. 추세(200dSMA)·변동성(실현)·breadth(>MA 비율)→레짐 상태.",
 "레짐을 *연속*으로 산출(이산 버킷 아님)+히스테리시스→ 코어/위성/분봉 예산·max gross·LETF dose·cash floor 결정.","멍청·사전등록·백테스트된 분류기만. 화려한 레짐모델=과적합 금지.")
layer(3,"Sleeve Library (고정 전략군)","① Core beta(SPY/QQQ/cash) ② Tactical trend+LETF ③ Cross-sectional alpha(재심사) ④ Intraday Lab.",
 "각 sleeve가 독립적으로 타깃 산출. 사전 빌드된 *고정* 집합(즉흥 추가 금지).","위성은 코어 대비 한계기여로만 평가. LETF=실제 3x 경로·dose only.")
layer(4,"Meta-Allocator (핵심·자본배분권)","일봉. 레짐→sleeve 연속 비중.",
 "비중 ∝ 추세강도×역변동성, 레짐 게이트. *고정비중(1/N) baseline을 OOS·비용후 못 이기면 폐기*.","walk-forward 검증된 규칙만. 분봉 lab의 요청 직접 안 받음. DD 한도 초과 risk-on 금지.")
layer(5,"Risk Engine","portfolio 일별 + intraday stop. 결정적 veto.",
 "DD ladder·vol-target·per-name cap·gap. capital-veto(live 차단) vs research-veto(연구 차단) 분리.","VRP류=capital-veto지 research-veto 아님(tail-hedge 변형 연구 허용).")
layer(6,"Execution","일봉 next-bar(신호 t→진입 t+1). 분봉=시뮬.",
 "스프레드/수수료/차입/충격/부분체결 비용 모델. 분봉은 슬리피지 보수적.","look-ahead 구조적 차단. 분봉 체결가정 현실화(원하는 봉서 못 채움).")
layer(7,"Attribution & Monitoring","P&L 분해 + 라이브 대시보드.",
 "core / 각 위성 한계기여 / 얼로케이터 타이밍(고정비중 대비) / 분봉 gross−비용−슬리피지 따로.","'누가 돈 벌었나' 항상 분해. health.json·kill-switch·heartbeat.")
layer(8,"Governance & Promotion","disposition 분류 + 사람 자본 게이트.",
 "live/paper/변형/sandbox/영구폐기. 자동개선=research candidate만(라이브 변경은 사람+동결 프로토콜).","irreducible 인간 게이트: 자본증액·코드변경·리스크한도·벤더변경·킬후재시작. trial registry.")

H(d,"C. 권한 구조 (누가 무엇을 결정하나)")
TBL(d,[["엔진","주기","권한"],
 ["Daily Regime Engine","장마감 1회","시장상태 판단·예산한도·max gross·LETF dose·cash floor 결정"],
 ["Capital Allocator","일봉","실제 자본 배분권 보유. 분봉 lab 요청 직접 안 받음"],
 ["Intraday Lab","프리마켓/장중","scanner/VWAP/ORB. Daily가 준 예산 내 paper만. 코어 자본 침범 금지"],
 ["Risk Engine","일별+장중","portfolio·intraday risk. capital-veto / research-veto 분리"],
 ["Attribution","일별","core·위성·allocator·intraday gross/비용 전부 분해"]])

H(d,"D. 빌드 순서 (빅뱅 금지)")
BUL(d,["1단계: Core(SPY/QQQ/cash) + 위성1(trend/LETF) + 멍청한 allocator(고정 70/25/5) + attribution. → 'core 위에 위성이 진짜 보태나' 확인.",
 "2단계: 연속 allocator(추세강도×역변동성) vs 고정비중 비교 — OOS서 이겨야 채택.",
 "3단계: 위성 하나씩 추가(cross-sectional alpha 재심사), 각자 한계기여 증명.",
 "4단계: Intraday Lab(처음부터 capital 배분 금지·paper만·gross−비용 분해).",
 "5단계: DL 보조(분봉 데이터 쌓인 뒤 메타라벨러)."])
P(d,"참고: 원래 PRAMANA charter(퀀트 딥리서치.docx)도 '일봉 swing/mid-long, 주1회~월1회 리밸런스 우선, intraday는 나중에 보조'라 적었다 — V2의 분봉-나중 순서와 일치.",i=True,c=GRAY)

H(d,"E. 회의 안건 (상충·미해결·고도화 필요)",c=AMBER)
P(d,"아래는 합의 전 — 다음 회의에서 결정/검증할 것.",b=True,c=AMBER)
TBL(d,[["#","안건","쟁점"],
 ["A1","목적함수 확정","'공격+낙폭한도(MDD≤30%)' vs '무제한 공격' vs '위험조정 우선' — 미정. 실제 운용규모·시계·감내낙폭 필요."],
 ["A2","Meta-Allocator 과적합","얼로케이터가 v2의 새 과적합 표면. '고정비중 OOS로 못 이기면 폐기' 규율을 어떻게 박을지(walk-forward·휩쏘비용·연속vs이산)."],
 ["A3","quality 부활 여부","원본 B3 재검토 필요. 최근절반 decay(IC-IR 0.046) 우려 — 코어 대비 한계 무상관 기여로 살릴지."],
 ["A4","기관신호 강도/소스","13F 45일 지연=느린 틸트만. 어느 기관·어떤 신호(13F/ETF flow)·복제규칙 사전등록(이것도 과적합 가능)."],
 ["A5","분봉 진지도","진짜 일중 단타인지 vs '일봉보다 자주 보는 스윙'인지 — 반대 강도가 갈림. 분봉=별도 비용·증거 프로젝트."],
 ["A6","숏/차입 현실성","L/S 절대수익 살리려면 차입비용·locate·Reg SHO 반영 필수(Codex: 0.5% 차입은 환상)."],
 ["A7","데이터 이원화","실자본 전 2 독립피드+브로커 대조 필요(무료 단일=레버 부적합). 분봉 데이터 벤더·비용 미정."]])
d.save(os.path.join(OUT,"PRAMANA_V4_Design.docx"))
print("✅ 2/3 PRAMANA_V4_Design.docx")

# ════════════════════════════════════════════════════════════════════
# DOC 3 — 마스터 다이제스트 (모든 연구 한 파일)
# ════════════════════════════════════════════════════════════════════
d=newdoc()
d.add_heading("PRAMANA — 마스터 다이제스트",0)
P(d,"전체 연구를 한 파일로 · 2026-06-11 · 가지고 다니는 정본 요약",i=True,c=GRAY)
P(d,"이 파일 하나면 PRAMANA 전체(목표·연구아크·실증결과·아키텍처·현재상태·파일위치)를 파악한다. 원본 100+ 문서를 다 읽을 필요 없음 — 필요하면 §6 인덱스로 찾아간다.",b=True)

H(d,"1. 프로젝트 한눈에")
BUL(d,["정체: 솔로+AI용 US 주식/ETF systematic 페이퍼 트레이딩·검증 OS (종목 찍는 AI 아님).",
 "목표: 데이터·비용 정직성 → 검증 → 무상관 전략 결합·공격적 사이징. 결정적 risk veto·LLM off-path·사람이 자본 게이트.",
 "현재 상태: 단순/선형 팩터 family 전멸(정직한 no-edge). 풀북은 SPY를 위험조정으로 못 이김. trend+LETF 엔진만 유망(검증중).",
 "다음: V2 core-satellite 재설계(SPY/QQQ 코어+공격 위성). forward 페이퍼 라이브 가동중(2026-06-09 인셉션)."])

H(d,"2. 연구 아크 — DR-1~4 적대 락 시퀀스")
P(d,"방법: ChatGPT primary → Claude deep-research red-team → 병합 lock. 알파 짜기 *전에* 인프라(벤치·유니버스·비용·PIT데이터·검증·모델맵)를 사실로 고정. 모든 수치 임계값은 원전 인용 또는 house-rule 라벨 강제.",i=True,c=GRAY)
TBL(d,[["라운드","스코프","Claude red-team 핵심","최종 락"],
 ["Stage 0A","축 고정(벤치/유니버스/비용/데이터)","권한 위상 고정(신호≠주문·risk veto·LLM off-path)","authority topology LOCK"],
 ["DR-1","4축 사실락","PATCH: survivorship/delisting(~−30%)가 알파 빌드 포트폴리오에 집중→선결","cost분해·PIT-before-alpha LOCK"],
 ["DR-1A","KR 무료소스 락","PATCH+US-first REJECT: KRX300 TR·거래세·OpenDART는 검색 false-negative(다 실재)","KR 7항목 상향(잠정)"],
 ["DR-2A","KR PIT 데이터","PATCH: PIT멤버십 'Lock now' 과신(2014플로어·free-float없음)·earnings는 무료(KIND/DART)","OpenDART LOCK·멤버십 Provisional"],
 ["DR-2B","US PIT 데이터","PATCH: 'CRSP delisting 지금 Lock' 깸(코드 −55/−66/−99·Shumway)→무료EDGAR+저가벤더 prototype","US-core first LOCK·CRSP DEFER"],
 ["DR-3","검증 프로토콜","PATCH: t>3=Harvey-Liu-Zhu(시변)·DSR는 N없이 계산불가·PBO 원전컷=0.05뿐","방법 LOCK·임계값 provisional"],
 ["DR-4","모델 후보맵","PATCH: 진짜 baseline=1/N(MVO아님)·penalized/tree=challenger·TSFM=reject-for-alpha","boring baseline first·deep BAN"]])
P(d,"전 라운드 PATCH. 가장 큰 교정: ① DR-2B가 'CRSP 지금 Lock'을 CRSP 자체 누락코드+문헌으로 깸 → 무료-우선 prototype. ② DR-4가 boring baseline을 1/N으로 낮추고 deep/TSFM을 alpha서 reject. 결론 시퀀스: 인프라 먼저, US-only, 1/N부터.",i=True,c=GRAY)

H(d,"3. 실증 결과 전체 (Phase 1A/1B + V2/V3)")
P(d,"전부 페이퍼·단발 사전등록·튜닝 없음·goalpost 결과 전 고정.",i=True,c=GRAY)
TBL(d,[["실험","핵심 수치","판정"],
 ["B0 broad (S&P500 PIT cap-weight)","실제 SPY corr 0.998·+0.55%p/yr","PASS(파이프 검증)"],
 ["B1 (1/N)","corr-SPY 0.937","PASS"],
 ["B2 value / B4 mom / B5 lowvol","|IC|<0.01 (전부)","DEAD"],
 ["B3 quality broad","IC-IR 0.220·net +4.41%","SURVIVE-screen"],
 ["B3 quality quarantine","decay 0.42→0.046·long-only<CW −1.15%","FAIL(TERMINATE)"],
 ["Phase1B 결합(blend/rank/ridge)","ridge OOS IC-IR 0.014·전부<CW","FAIL(family 종료)"],
 ["US event-drift","no-decay·net<CW −0.90%","FAIL"],
 ["US small/mid blend","t≈3.75나 유동성집중→liquid 슬라이스 음수","FAIL"],
 ["V2 True PEAD","IC −0.0153(부호역전)","REJECT"],
 ["V2.1 analyst revision","구조 GO·가격 HOLD(2018 무료샘플만)","HOLD(데이터축)"],
 ["V3 trend overlay","Sharpe 0.85·corr +0.018","SURVIVE"],
 ["V3 LETF","combo +1.21%/yr(매칭DD)","SURVIVE"],
 ["V3 VRP short-vol","최악월 −62%·MDD −92%","REJECT"],
 ["V3 reversal","gross Sharpe −0.52·turnover 937%","REJECT"]])
P(d,"멀티시점 net(2026-06): 풀북 3M/6M/1Y/풀사이클 전부 SPY에 위험조정으로 짐. 풀사이클 3x +699% vs SPY +307%=레버(MDD −45% vs −34%·Sharpe 0.82<0.90). trend+LETF만 ~400일 +22% vs SPY +15%·Sharpe 1.38(가설).",b=True,c=AMBER)
P(d,"메타패턴: 모든 횡단면 신호가 1/N은 이기나 cap-weight long-only엔 진다=2016-26 메가캡 지배. 신호는 small/mid에 더 살지만 유동성+비용이 거래가능분을 죽인다.",i=True,c=GRAY)

H(d,"4. 아키텍처 + 잠금 결정 + Codex 리뷰")
P(d,"8 레이어: 데이터→신호→앙상블→포트폴리오→리스크엔진→실행→모니터/kill→거버넌스 (모듈 data/features/combined_book/risk_engine/execution/monitor/book_runner). 전부 config-driven·API-free·동결재현.")
P(d,"잠금 결정: US-only·가상 ₩100M(no live)·Sharadar 데이터·모델공간 CONSTRAINED(정규화선형+GBM challenger·deep BAN)·결정적 risk veto·LLM off-path·사람 자본 게이트·사전등록 kill·data-gate-before-model-gate·동결재현.",i=True,c=GRAY)
P(d,"Codex 적대 리뷰 = REVISE(not fundable as-is): '인-샘플 equity 베타 타이밍을 분산북으로 포장'·forward Sharpe 0.35-0.55·무상관 +0.06은 스트레스서 +0.25~0.60·same-close 누수(→next-bar 수정)·차입 0.5%=환상·'지금 $75k 넣겠나? NO'. 수용: next-bar 재측정·현실 비용·forward 6-12개월 후 사람 게이트.",c=RED)

H(d,"5. v1–v3 → V4 전환")
P(d,"진단: 기계는 정확(corr 0.998)·단순팩터 no-edge 정직 입증. 결함=잣대 미스칼리브레이션(기관용)·이진 폐기·core 부재(SPY 못 이김)·analysis-paralysis. 상세 → PRAMANA_PostMortem_v1to3.docx.")
P(d,"V4 방향: core-satellite(SPY/QQQ 코어+공격 위성)·전략군 고정 비중 레짐자동·리스크엔진 판정언어 교체(capital/research veto 분리)·DL=필터·분봉 최강격리·attribution 필수·빅뱅금지. 상세 → PRAMANA_V4_Design.docx.")

H(d,"6. 전체 파일 인덱스 (어디서 무엇을 찾나)")
TBL(d,[["분류","위치","내용"],
 ["연구 아크(DR)","pramana_research/00~15","DR-1~4 chain·primary+audit·Phase1A 설계"],
 ["실증 결과","phase1a/reports/","B0~V3 실험 결과 정본"],
 ["설계·잠금","integrated/","8레이어·V3 설계·락시트·Codex handoff·기술보고서"],
 ["DR 폴더 고유","Desktop/DR/","ChatGPT 원본 PDF·Stage0A locks·퀀트딥리서치 charter(L0-L12)"],
 ["통합 아카이브","PRAMANA_ARCHIVE/","위 전부를 한 곳에 dedup·분류(00_INDEX.md 참조)"],
 ["V4 산출물","PRAMANA_V4/","이 다이제스트·v1-v3 사후분석·V4 설계도"],
 ["코드","phase1a/engine/","27 모듈·forward_runner·multi_anchor_sim·dashboard"],
 ["마일스톤","phase1a/registry/","phase1a_milestones.csv (54개)"]])

H(d,"7. 현재 운영 (라이브)")
BUL(d,["forward_runner.py — 무인 일1회 trend+LETF 페이퍼(무료 EOD·fail-closed). 라이브 인셉션 2026-06-09.",
 "outputs/forward_dashboard.html — 라이브 대시보드(1h refresh).",
 "outputs/multi_anchor.html — 멀티시점 vs SPY.",
 "GPU·클라우드 불필요(ML 미사용·CPU pandas). 미니PC cron `0 6 * * 2-6`."])
d.save(os.path.join(OUT,"PRAMANA_MASTER_DOSSIER.docx"))
print("✅ 3/3 PRAMANA_MASTER_DOSSIER.docx")
print(f"\n→ {OUT}/")
for f in sorted(os.listdir(OUT)):
    if f.endswith(".docx"): print(f"   {f}  ({os.path.getsize(os.path.join(OUT,f))//1024}KB)")
