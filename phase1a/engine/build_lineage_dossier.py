#!/usr/bin/env python3
"""PRAMANA 계보 Dossier (DOCX) — v1~V7 진화·결과물·문제점·트러블슈팅 + 12/6/3개월/풀 멀티앵커 표(실계산).
free yfinance·비용후. PAPER·RESEARCH_ONLY. python engine/build_lineage_dossier.py → PRAMANA_V4/PRAMANA_Lineage_Dossier.docx"""
import os, warnings; warnings.filterwarnings("ignore")
import numpy as np, pandas as pd, yfinance as yf
from docx import Document; from docx.shared import Pt, RGBColor, Inches; from docx.enum.text import WD_ALIGN_PARAGRAPH
HERE=os.path.dirname(os.path.abspath(__file__)); ROOT=os.path.dirname(HERE); REPO=os.path.dirname(ROOT)
OUT=os.path.join(REPO,"PRAMANA_V4","PRAMANA_Lineage_Dossier.docx")
CAP=1e8; RF=0.05; TVOL=0.28
# ── 멀티앵커 실계산 (V4~V7 동일 엔진·비용후) ──
px=yf.download(["SPY","QQQ","DBMF","GLD","IEF"],period="max",interval="1d",auto_adjust=True,progress=False)
px=(px["Close"] if isinstance(px.columns,pd.MultiIndex) else px).dropna()
ret=px.pct_change(); core=0.5*ret["SPY"]+0.5*ret["QQQ"]; rvol=core.rolling(20).std()*np.sqrt(252)
days0=core.dropna().index; days0=days0[days0>=rvol.dropna().index[0]]
def ddscale(dd): return 0.2 if dd<=-0.40 else 0.4 if dd<=-0.30 else 0.7 if dd<=-0.20 else 1.0
def v4(days):  # Core Beta 1.0x (SPY/QQQ 50/50·무레버)
    return CAP*(1+core.reindex(days).fillna(0)).cumprod()
def lev_book(days, cap, mf):  # V5(mf=0·cap2.0)·V6(mf=0.15·cap1.5)
    nav=[CAP]; peak=CAP
    for i in range(1,len(days)):
        d,p=days[i],days[i-1]; rv=rvol.get(p,np.nan)
        L=min(cap,TVOL/rv) if (pd.notna(rv) and rv>0) else 1.0; L*=ddscale(nav[-1]/peak-1)
        r=(1-mf)*(L*core.get(d,0)-(L-1)*RF/252)+mf*ret["DBMF"].get(d,0)
        nav.append(nav[-1]*(1+r)); peak=max(peak,nav[-1])
    return pd.Series(nav,index=days)
W7={"SPY":0.25,"QQQ":0.25,"DBMF":0.25,"GLD":0.15,"IEF":0.10}
def v7(days):  # 4-sleeve 1.0x
    b=sum(W7[t]*ret[t] for t in W7); return CAP*(1+b.reindex(days).fillna(0)).cumprod()
def bh(s,days): return CAP*(1+ret[s].reindex(days).fillna(0)).cumprod()
def stat(nav):
    r=nav.pct_change().dropna(); return (nav.iloc[-1]/nav.iloc[0]-1, (nav/nav.cummax()-1).min(), (r.mean()/r.std()*np.sqrt(252) if r.std()>0 else float('nan')))
end=days0[-1]
ANCH=[("3개월 전 진입",91),("6개월 전 진입",182),("12개월 전 진입",365),(f"풀사이클 ({days0[0].date()}~)",99999)]
BOOKS=[("V4 Core Beta 1.0x",lambda d:v4(d)),("V5 Leveraged cap2.0",lambda d:lev_book(d,2.0,0.0)),
       ("V6 Diversified cap1.5+DBMF",lambda d:lev_book(d,1.5,0.15)),("V7 4-sleeve",lambda d:v7(d)),
       ("QQQ",lambda d:bh("QQQ",d)),("SPY",lambda d:bh("SPY",d))]
TABLE={}
for an,dd in ANCH:
    start=end-pd.Timedelta(days=dd); sl=days0[days0>=start] if dd<99999 else days0
    TABLE[an]={nm:stat(fn(sl)) for nm,fn in BOOKS}
print("멀티앵커 실계산 (누적%/MDD%/Sharpe):")
for an in TABLE:
    print(f"\n[{an}]"); [print(f"  {nm:<28}{a[0]*100:>7.1f}{a[1]*100:>7.0f}{a[2]:>7.2f}") for nm,a in TABLE[an].items()]

# ── DOCX ──
doc=Document()
def H(t,l=1,c="1F3864"):
    h=doc.add_heading(t,level=l)
    for r in h.runs: r.font.color.rgb=RGBColor.from_string(c)
    return h
def P(t,b=False,sz=10,c=None,it=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=b; r.italic=it; r.font.size=Pt(sz)
    if c: r.font.color.rgb=RGBColor.from_string(c)
    return p
def bullet(t,sz=9.5):
    p=doc.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(sz); return p

ti=doc.add_heading("PRAMANA — 전략 계보 Dossier",0)
P("v1~v3 단순팩터 → V4 Core Beta → V5 Leveraged → V6 Diversified → V7 4-sleeve",True,12,"C55A11")
P("PAPER only · NO LIVE · RESEARCH_ONLY · 솔로+AI · US 주식/ETF · 가상 ₩100M · 2026-06-12",False,9,"808080",True)
P("이 문서: 각 세대가 무엇을 만들었고(결과물) 무엇이 문제였고(한계) 왜 다음으로 진화했나 + 12/6/3개월·풀 멀티앵커 성과표(실계산) + 트러블슈팅 종합. 한눈에 'V7까지 왜·어떻게'.",False,9.5,it=True)

H("0. 한눈에 — 메타 결론",1)
for t in ["검증되어 Production에 올릴 수 있는 알파는 아직 0: v1~V7 전부 베타·레버·분산·리스크통제지 *예측 스킬* 아님('쉬운 알파 없음' 7세대 반복 실증). 단 Alpha Lab은 별도 paper 연구(자본권한 0)로 유지 — v2 event-driven momentum forward 관찰 중.",
          "각 피벗은 이전의 *정직한 실패*에서 나옴 — 가짜 승리를 안 만들고 매번 '이건 알파 아님'을 인정하며 진화.",
          "V7 4-sleeve = 알파가 아니라 *목적함수 선택*: **크래시 생존 목적함수에선 가장 설득력(Sharpe·MDD 최선)·최대복리 목적함수에선 QQQ 열위**(수익 ~절반 양보=보험료). 검증은 forward 12개월.",
          "V5→V7은 파레토 개선이 아니라 효율적 프론티어 *위에서의 이동*(위험↔수익 트레이드) — Sharpe 개선조차 분산 프리미엄이지 예측력 아님 = 알파 못 만든 증거.",
          "시스템의 진짜 가치 = *가짜 알파를 거부한 정직성*(단순팩터·레버·LETF·throttle·core-switching·RVOL look-ahead 누수까지 잡아 죽임)."]:
    bullet(t)

H("1. 불변 LOCK 시트 — 버전이 바뀌어도 변하지 않는 것",1)
P("v1~V7 전 세대를 관통하는 헌법. 버전업(피벗)은 이 LOCK *위에서만* 일어난다. 정본: docs/context/LOCKS.md · PRAMANA_V4_LOCK_SHEET_v0.4.md",False,9,"808080",True)
P("A. 거버넌스 (불변)",True,10,"2E5496")
for t in ["PAPER only · NO LIVE — 실자본은 Promotion Gates(crash-pack + 12mo forward + 2-feed reconciliation + attribution + 사람 게이트) 전부 통과 후만",
          "US only — 미국 시장 포커싱(crypto 제외)",
          "core-satellite — SPY를 버리고 이기려 하지 않음(싸우지 말고 깔고 얹기)",
          "목적함수 = max 초과수익 s.t. MDD ≤ X (수익극대화 + 파산 회피)",
          "사람이 자본 게이트 — 자본증액/코드변경/리스크한도/벤더변경/킬후재시작은 사람 승인",
          "검증되어 Production 승격 가능한 알파 = 아직 0 (연구는 자본권한 0으로 OPEN)"]: bullet(t)
P("B. 방법론 (불변)",True,10,"2E5496")
for t in ["next-bar 실행 — same-close 체결 금지·look-ahead 차단",
          "attribution 필수 — 어느 sleeve가 기여했나 분해",
          "결정적 risk engine 최종 veto",
          "baseline 우선 — 복잡 모델·allocator는 1/N·고정비중을 OOS 비용후 이겨야 채택",
          "LLM · TSFM 직접 트레이딩 금지 (off-path 보조만)",
          "production = 일봉 · 데이터 무결성 = 실자본 전 2 독립피드",
          "사전등록 kill — 결과 본 뒤 goalpost 이동 금지(= config-mining 금지)",
          "Macro Emergency Override — 위기 확정 시(QQQ 200일선 하회·고점 −25%·VIX 급등·breadth 붕괴·credit·연속 gap-down) *공격 파트만*(레버·V8 shadow·Alpha Lab 신규 risk-on) 차단 + 사람 리뷰 트리거. **코어 4-sleeve 1.0x는 면역**(자동 코어전환 아님 = regime-switch/throttle 기각 교훈)·후행이라 첫 −25%는 못 피함(예측 아니라 *추가 위험 차단 + 사람 게이트*)"]: bullet(t)
P("C. 8레이어 아키텍처 (불변 골격)",True,10,"2E5496")
lt=doc.add_table(rows=1,cols=2); lt.style="Light Grid Accent 1"
hc=lt.rows[0].cells; [setattr(hc[i].paragraphs[0].add_run(x),"bold",True) for i,x in enumerate(["레이어","역할"])]
for a,b in [("L1 데이터","Sharadar 일봉/PIT/펀더멘털 + yfinance forward"),
            ("L2 신호/레짐","후행 지표(20일 변동성·낙폭) — 예측 아님"),
            ("L3 북(sleeve)","4-sleeve(주식·MF·금·채권) 등 sleeve 구성"),
            ("L4 배분(allocator)","고정비중 우선(적응형은 baseline 이겨야)"),
            ("L5 리스크엔진","결정적 veto·vol-target·DD ladder"),
            ("L6 실행","next-bar·비용/financing 모델"),
            ("L7 귀속/모니터","attribution·forward 판정표·대시보드"),
            ("L8 거버넌스","paper·사람 자본게이트·12mo STOP·사전등록 kill")]:
    r=lt.add_row().cells; r[0].text=a; r[1].text=b
    for cc in r:
        for pp in cc.paragraphs:
            for rr in pp.runs: rr.font.size=Pt(9)
P("D. 데이터 역할 분담 (불변)",True,10,"2E5496")
for t in ["Sharadar (유료) = 일봉/PIT/펀더멘털/survivorship-safe *backtest backbone* — 알파를 *만든* 게 아니라 *가짜 알파를 걸러준* 데이터(self-built S&P500 PIT가 실제 SPY와 corr 0.998).",
          "yfinance (무료) = forward/probe — intraday 5m는 60일 제약·proof 약함(FAST_REPLAY, NOT_PROOF).",
          "intraday Alpha Lab = 별도 벤더(Polygon/Alpaca/QuantRocket) 필요할 수 있음 — VWAP/ORB/quote는 Sharadar 본업 아님.",
          "결제 판단(불변 기준): ① v2가 무가치면 유료 보류 ② 후보는 잡히나 체결·스프레드·기간 제약으로 판단불가면 1개월 파일럿 ③ 목적 = '알파 구매'가 아니라 '데이터 병목 제거'."]: bullet(t)
P("E. 위기 데이터 처리 (불변)",True,10,"2E5496")
for t in ["닷컴(2000-02)·2008 같은 큰 위기 = 학습/파라미터 튜닝엔 *제외 또는 별도 분리*(과최적화 노이즈)·promotion stress 검증엔 **필수 포함**.",
          "'사람이 알아차리고 뺄 수 있다'는 가정 금지 — 위기는 실시간엔 단순 조정처럼 오고 −20/−30% 후에야 확정. 회피는 감이 아니라 Macro Emergency Override(시스템화·위 B)로만 인정.",
          "장기 stress proxy 한계 명시(2000 managed-futures sleeve = cash 부재 → 4-sleeve에 불리)·'닷컴 때문에 전부 폐기'도 '닷컴 무시 통과'도 금지 → UNKNOWN 라벨."]: bullet(t)
P("→ 이 LOCK들은 버전(V8·V9…)이 와도 안 바뀐다. 바뀌면 그건 PRAMANA가 아니라 다른 프로젝트다.",True,9.5,"C00000")

H("2. Phase A — 단순팩터 시대 (v1~v3): '쉬운 알파 없음' 실증",1)
P("결과물 (= 연구 자산):",True,10)
for t in ["데이터 파이프라인 검증: self-built S&P500 PIT cap-weight가 실제 SPY와 corr 0.998 (survivorship/PIT 정확).",
          "팩터 6 family(value·momentum·quality·lowvol·event) IC 측정 + 결합(ridge/GBM) + 풀북.",
          "정직한 negative 리포트 — 시스템이 가짜 알파를 안 만든 게 자산."]: bullet(t)
P("드러난 문제점:",True,10)
for t in ["단순/선형 팩터 전부 비용후 SPY/QQQ 못 이김 · Rank IC≈0 · quality만 약하게 생존하다 식음(IC-IR 0.22→0.046).",
          "결합(blend/ridge)도 OOS cap-weight 못 이김 · v3 풀북도 SPY를 위험조정으로 못 이김(레버지 엣지 아님)."]: bullet(t)
P("트러블슈팅: same-close 체결 누수 발견 → next-bar 실행으로 수정(Codex REVISE) · survivorship → self-built PIT.",False,9.5,"C00000")
P("→ 판정: US 단순/선형 팩터 family TERMINATE. → 왜 피벗: 'SPY와 싸워 못 이긴다 → 싸우지 말고 깔고 얹자' = core-satellite (V4로 명명, 내부 v1/v2/v3와 혼동 방지).",True,9.5)

H("3. Phase B — core-satellite·레버·분산 시대 (V4~V7)",1)
GEN=[("V4 — Core Beta Forward Book","SPY/QQQ 1.0x 베타북 (레버 격리)",
      ["production_book.py(target JSON) · forward_runner.py(무인 일1회·라이브 2026-06-09)"],
      ["QQQ를 못 넘음(설계상·베타북) · −35%MDD에 레버 맞추기 = 과거 낙폭에 끼워맞추는 'backward knob'(Codex: PRODUCTION_UNSAFE)"],
      "레버를 격리 sleeve로 빼고 기본 1.0x · overlay −0.14%=노이즈 OFF",
      "3M sim: Core Beta +10.9% > SPY +7.6% but QQQ +14.3%엔 짐 → 레버로 넘어보자 (V5)"),
     ("V5 — Aggressive Leveraged Core Beta","vol-target 28%·캡 2.0x·DD ladder",
      ["aggressive_book.py · forward_runner_v5.py(라이브 2026-06-11) · forward 정량 판정표 · multi_anchor_v5"],
      ["레버지 알파 아님(Sharpe≈QQQ) · SPY/QQQ와 *같이* 낙폭(6mo −20%) · forward −70%+ 가능(benign 샘플·2008/닷컴 없음) · no-ruin 아님(gap 무력)"],
      "vol-target procyclical 발견(저점 디레버·반등 놓침 −6.3%p) · 판정표 라이브-only·상방참여 60% 미달 정직 표기",
      "'같이 낙폭'이 핵심 문제 → 구조적 분산 (V6)"),
     ("V6 — Diversified Aggressive","85% 레버드 Core(cap1.5) + 15% managed-futures(DBMF)",
      ["v6_book.py · forward_runner_v6.py · derisk_diversify_test.py(DBMF 2022 +21.5% 방어)"],
      ["alpha 아니라 *보험료*(상승장 수익↓·2023 드래그) · 여전히 equity-dominant(85% 주식) · DBMF 짧은 역사(2019~·2008 없음)·MF lost-decade(2011-20)·샘플 의존"],
      "판정표 라이브-only · no-ruin docstring 제거 · 대시보드 in-sample 프레이밍(Codex 전체검수 fix)",
      "equity-dominant라 큰 크래시 여전 같이 맞음·'왜 QQQ 대신?' 약함 → 진짜 분산 4-sleeve (V7)"),
     ("V7 — 4-sleeve Paper Core Candidate (현재)","Eq50(SPY/QQQ)+MF25(DBMF)+Gold15(GLD)+Bond10(IEF) + Risk Throttle + Alpha Lab + Risk Monitor",
      ["forward_runner_v7.py(4-sleeve 1.0x·라이브 2026-06-11) · Risk Monitor · crash_pack_throttle · Alpha Lab v0(인프라)/v1(DEAD·forward 관찰)/v2(event-driven 로그)"],
      ["QQQ bull 수익 ~절반 포기(목적함수 선택) · throttle 기각(crash-pack서 static 4-sleeve 못 이김) · Alpha Lab v1 setup DEAD(RVOL 누수·강세장 베타) · forward 0개월 · DBMF/GLD 짧은 역사 · ★ 닷컴류 stress: 2019~ MDD −18%는 *benign 샘플*(닷컴/2008 없음)·진짜 닷컴류엔 V7 1.0x도 −39%(2000 MF proxy 부재로 보수적·불완전)"],
      "코어 regime-switch 휩쏘 기각(코어 안 갈아탐) · throttle crash-pack 기각(대시보드 전용) · RVOL look-ahead 누수 수정 · 'Production Core'→'Paper Core Candidate'(Codex)",
      "(현재) forward 12개월 관찰 + Alpha Lab event-driven paper 로그·닷컴 caveat 추가"),
     ("V8 Candidate — Levered 4-sleeve (UNKNOWN · RESEARCH_ONLY · shadow only)","V7 분산북을 조심스럽게 레버(1.10~1.25x) · 자본권한 0 · 알파 아님",
      ["v8_levered_crashpack.py(grid crash-pack 1.0~1.5x) · V8_Candidate_Levered4sleeve_Protocol.md"],
      ["알파 아님 = 분산 프리미엄 레버 = 목적함수 이동(상승참여↑·하락방어) · 닷컴 proxy서 1.25x −49%(단 2000 MF=cash 부재) vs 2008 −32% 통과 = **UNKNOWN** · 1.35x/1.5x 폐기(2008서도 −35%) · forward 레버 꼬리(V5 −70% 교훈)"],
      "Codex가 equity proxy 결함(VFINX≠SPY/QQQ·QQQ 닷컴 −78% 누락) 잡음 → 실 SPY/QQQ로 수정 → 판정 UNKNOWN(전멸도 통과도 아님)",
      "1.10~1.25x 자본0 shadow forward로 'UNKNOWN 추적'(승격도 폐기도 아님)·실자본 금지·옵션 컨벡시티는 나중 후보")]
for nm,sub,res,prob,trb,nxt in GEN:
    H(nm,2,"2E5496"); P(sub,True,10,"C55A11")
    P("결과물:",True,9.5); [bullet(t) for t in res]
    P("문제점/한계:",True,9.5); [bullet(t) for t in prob]
    P(f"트러블슈팅: {trb}",False,9.5,"C00000")
    P(f"→ 왜 다음으로: {nxt}",True,9.5)

H("4. 멀티앵커 성과표 (12/6/3개월·풀 · 비용후 · 실계산)",1)
P("동일 엔진으로 각 진입 시점부터 현재까지. DBMF(2019~) 제약으로 풀사이클은 2019년부터. 누적·MDD·Sharpe. *in-sample backtest지 forward 판정 아님.*",False,9,"808080",True)
for an in TABLE:
    P(an,True,10,"2E5496")
    tb=doc.add_table(rows=1,cols=4); tb.style="Light Grid Accent 1"
    hd=tb.rows[0].cells; [setattr(hd[i].paragraphs[0].add_run(x),"bold",True) for i,x in enumerate(["북","누적 %","MDD %","Sharpe"])]
    for nm,a in TABLE[an].items():
        c=tb.add_row().cells; c[0].text=nm; c[1].text=f"{a[0]*100:+.1f}"; c[2].text=f"{a[1]*100:.0f}"; c[3].text=f"{a[2]:.2f}"
    P("",False,4)
P("읽는 법: V5는 레버로 누적↑지만 MDD도↑·Sharpe는 QQQ급(레버지 알파 아님). V6/V7은 MDD↓·Sharpe↑지만 상승장 누적은 QQQ에 양보(분산=보험료). **크래시 생존 목적함수에선 V7 4-sleeve가 최선(Sharpe·MDD)·최대복리 목적함수에선 QQQ 우월(누적). 어느 게 '낫다'가 아니라 어느 *목적*이냐의 문제.** V5→V7은 프론티어 위 이동이지 돌파(알파) 아님.",False,9.5,it=True)

H("5. 트러블슈팅 종합 (전 세대 버그/교정 — 정직성 기록)",1)
tt=doc.add_table(rows=1,cols=3); tt.style="Light Grid Accent 1"
h=tt.rows[0].cells; [setattr(h[i].paragraphs[0].add_run(x),"bold",True) for i,x in enumerate(["세대","문제","교정"])]
for a,b,c in [("v3","same-close 체결 누수","신호 t→진입 t+1 (next-bar)"),
              ("v1~v3","survivorship/PIT 누수","self-built S&P500 PIT (corr 0.998)"),
              ("V4","레버 = backward knob","레버 격리 sleeve·기본 1.0x"),
              ("V5","vol-target procyclical(저점 디레버)","발견·정직 표기·보험료로 수용"),
              ("V5/V6","판정표 수익-only 합격·no-ruin 과신","판정표 라이브-only·no-ruin 제거"),
              ("V7","코어 regime-switch = 마켓타이밍","휩쏘 기각·코어 영구 고정"),
              ("V7","throttle 비대칭(반등 못 켬)","crash-pack 기각·대시보드 전용"),
              ("Alpha Lab v1","RVOL look-ahead 누수","entry-time 누적 RVOL로 수정"),
              ("인프라","커밋 email 차단·무료 2nd feed 막힘","click6067 email 고정·2-feed=유료 대기")]:
    r=tt.add_row().cells; r[0].text=a; r[1].text=b; r[2].text=c
    for cc in r:
        for pp in cc.paragraphs:
            for rr in pp.runs: rr.font.size=Pt(9)

H("6. 현재 위치 / 다음 (3트랙만 — 능동 개발 늘리면 산으로 감)",1)
P("Track 1 — Core Forward + V8 shadow",True,10,"2E5496")
for t in ["순수 4-sleeve(V7) 1.0x 12개월 paper forward: QQQ·60/40 대비 누적/MDD/회복/Sharpe·위기방어. 절대 건드리지 말 것: LETF·throttle·core switching(기각됨).",
          "V8 Candidate 1.10~1.25x = 자본0 shadow forward로 'UNKNOWN 추적'(승격도 폐기도 아님·코어 1.0x와 분리·1.35x+ 폐기)."]: bullet(t)
P("Track 2 — Alpha Lab + LLM catalyst",True,10,"2E5496")
for t in ["v1 = DEAD 기록 + 원형 forward 관찰 · v2 = event-driven momentum forward 로그 · 튜닝 금지.",
          "LLM catalyst 분류 1회 = v2 후보를 실제 뉴스로 A/B/C/D 재분류(proxy C급이 진짜 이유 없는지·매수결정권 0)·또 테마베타면 Alpha Lab 저우선 강등.",
          "TSFM/Chronos = *나중* meta-labeler challenger(off-path·내일수익 예측 X·setup 메타라벨/tail-risk 후보·얕은 모델보다 OOS 비용후 좋아야 유지)."]: bullet(t)
P("Track 3 — 데이터 판단",True,10,"2E5496")
for t in ["Sharadar = 계속 daily backbone(universe/유동성/상폐/gap후보/earnings 맥락에도 재활용).",
          "분봉 = yfinance fast replay 가능하나 proof 약함(NOT_PROOF).",
          "유료 intraday = v2가 ① 무가치면 보류 ② 후보는 잡히나 체결/스프레드/기간 제약으로 판단불가면 1개월 파일럿(목적='데이터 병목 제거'지 '알파 구매' 아님)."]: bullet(t)
P("금지(불변): v1 setup 튜닝 · Production 자본 투입 · LIVE. 전부 PAPER·자본권한 0.",True,9.5,"C00000")
P("한 줄: PRAMANA는 '알파를 못 찾은 실패 프로젝트'가 아니라, 가짜 알파를 계속 제거한 끝에 V7 4-sleeve core + 별도 Alpha Lab(paper)으로 분리된 paper research OS다.",True,10,"C55A11")
P("계보: v1~v3 단순팩터(SPY 못 이김·파이프 검증) → V4 베타북 → V5 레버(같이 낙폭) → V6 분산(보험료) → V7 4-sleeve(크래시 생존 선택) + Alpha Lab(급등주 paper 관찰) → V8 Candidate(분산북 레버·UNKNOWN shadow).",False,9.5,it=True)
P("정합(2026-06-12): 닷컴은 학습 노이즈로 제외 가능하나 *리스크 검증 제외 금지*(감 아니라 Macro Emergency Override로 시스템화) · V8은 확정도 폐기도 아닌 UNKNOWN shadow · V7은 닷컴 caveat(−39%) 추가 · 알파는 LLM catalyst 붙인 Alpha Lab에서 계속.",True,9.5,"C55A11")

doc.save(OUT)
print(f"\n✅ DOCX 생성: {OUT}")
print(f"   섹션: 메타결론·Phase A(v1~3)·Phase B(V4~V7 각 결과물/문제/진화)·멀티앵커표(12/6/3/풀)·트러블슈팅·현재/다음")
if __name__=="__main__": pass
